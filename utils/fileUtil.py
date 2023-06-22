# author:孔德昱
# datetime:2023/6/11 15:08
"""
description：处理word文档
"""

from docx import Document


def read_docx_file(file_path):
    document = Document(file_path)
    return document


def convert_document_to_string(doc):
    paragraphs = []
    for paragraph in doc.paragraphs:
        paragraphs.append(paragraph.text)
    return '\n'.join(paragraphs)


def convert_docx_to_txt(docx_file, txt_file):
    """
        将word文档转换为文本文档
        该函数只会将Word文档中的文本内容提取为纯文本，不包括格式、图像等其他元素。
    """
    # 打开Word文档
    doc = Document(docx_file)

    # 逐段读取文本内容并写入txt文件
    with open(txt_file, 'w', encoding='utf-8') as file:
        for paragraph in doc.paragraphs:
            file.write(paragraph.text + '\n')


def read_txt_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    return content


if __name__ == '__main__':
    convert_docx_to_txt("../input/test01.docx", '../input/test01.txt')
